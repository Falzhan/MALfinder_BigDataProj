import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import os

class AnimeDescriptor:
    def __init__(self):
        self.numeric_cols = ['Episodes', 'Score', 'Popularity']
        self.categorical_cols = ['Type', 'Genres', 'Themes', 'Demographics']
    
    def clean_string_list(self, text):
        """Clean strings containing lists"""
        if isinstance(text, str):
            return text.strip('[]').replace("'", "").split(', ')
        return []
    
    def generate_summary(self, df):
        """Generate analysis summary"""
        summary = []
        summary.append(f"Analysis based on {len(df)} anime results.")
        
        if 'Score' in df.columns:
            summary.append(f"Average score: {df['Score'].mean():.2f} (range: {df['Score'].min():.2f}-{df['Score'].max():.2f})")
        
        if 'Type' in df.columns:
            top_type = df['Type'].mode().iloc[0] if not df['Type'].empty else "N/A"
            type_count = df['Type'].value_counts().iloc[0] if not df['Type'].empty else 0
            summary.append(f"Most common type: {top_type} ({type_count} titles)")
        
        if 'Genres' in df.columns:
            df['Genres_clean'] = df['Genres'].apply(self.clean_string_list)
            all_genres = [g for genres in df['Genres_clean'] for g in genres if g]
            if all_genres:
                top_genre = pd.Series(all_genres).mode().iloc[0]
                genre_count = pd.Series(all_genres).value_counts().iloc[0]
                summary.append(f"Most frequent genre: {top_genre} (appears in {genre_count} titles)")
        
        return "\n".join(summary)
    
    def get_basic_stats(self, df):
        """Calculate basic statistics for numeric columns"""
        available_numeric = [col for col in self.numeric_cols if col in df.columns]
        numeric_stats = df[available_numeric].describe().round(2)
        
        categorical_stats = {}
        for col in self.categorical_cols:
            if col in df.columns:
                if col in ['Genres', 'Themes']:
                    clean_data = df[col].apply(self.clean_string_list)
                    counts = clean_data.explode().value_counts().head()
                else:
                    counts = df[col].value_counts().head()
                categorical_stats[col] = counts
        
        return {'numeric': numeric_stats, 'counts': categorical_stats}
    
    def generate_plots(self, df):
        """Generate visualization plots"""
        plots = {}
        
        # Score Distribution
        if 'Score' in df.columns:
            plt.figure(figsize=(12, 6))
            sorted_df = df.sort_values('Score', ascending=True)
            plt.plot(range(len(sorted_df)), sorted_df['Score'], marker='o')
            plt.grid(True, alpha=0.3)
            plt.xlabel('Rank')
            plt.ylabel('Score')
            plt.title('Anime Scores Distribution')
            plt.tight_layout()
            plots['score_dist'] = plt.gcf()
            plt.close()

        # Type Distribution
        if 'Type' in df.columns:
            plt.figure(figsize=(10, 6))
            type_counts = df['Type'].value_counts()
            type_counts.plot(kind='bar')
            plt.title('Distribution by Type')
            plt.xlabel('Type')
            plt.ylabel('Count')
            plt.xticks(rotation=45)
            plt.tight_layout()
            plots['type_dist'] = plt.gcf()
            plt.close()

        # Genre Distribution
        if 'Genres' in df.columns:
            plt.figure(figsize=(12, 6))
            genres = df['Genres'].str.split(',').explode()
            genre_counts = genres.value_counts().head(10)
            genre_counts.plot(kind='bar')
            plt.title('Top 10 Genres Distribution')
            plt.xlabel('Genre')
            plt.ylabel('Count')
            plt.xticks(rotation=45)
            plt.tight_layout()
            plots['genre_dist'] = plt.gcf()
            plt.close()

        # Theme Distribution
        if 'Themes' in df.columns:
            plt.figure(figsize=(12, 6))
            themes = df['Themes'].str.split(',').explode()
            theme_counts = themes.value_counts().head(10)
            theme_counts.plot(kind='bar')
            plt.title('Top 10 Themes Distribution')
            plt.xlabel('Theme')
            plt.ylabel('Count')
            plt.xticks(rotation=45)
            plt.tight_layout()
            plots['theme_dist'] = plt.gcf()
            plt.close()
        
        return plots
    
    def export_report(self, df, query, output_path='anime_analysis.docx'):
        """Export analysis to Word document"""
        doc = Document()
        
        # Title and metadata
        doc.add_heading('Anime Search Results Analysis', 0)
        doc.add_paragraph(f"Search Query: '{query}'")
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Summary
        doc.add_heading('Analysis Summary', level=1)
        doc.add_paragraph(self.generate_summary(df))
        
        # Basic Statistics
        doc.add_heading('Detailed Statistics', level=1)
        stats = self.get_basic_stats(df)
        
        # Numeric stats
        if not stats['numeric'].empty:
            doc.add_heading('Numeric Features', level=2)
            numeric_table = doc.add_table(rows=1, cols=len(stats['numeric'].columns) + 1)
            headers = numeric_table.rows[0].cells
            headers[0].text = 'Metric'
            for idx, col in enumerate(stats['numeric'].columns):
                headers[idx + 1].text = col
            
            for idx, row in stats['numeric'].iterrows():
                cells = numeric_table.add_row().cells
                cells[0].text = str(idx)
                for col_idx, value in enumerate(row):
                    cells[col_idx + 1].text = f"{value:.2f}"
        
        # Categorical stats
        doc.add_heading('Categorical Features', level=2)
        for category, counts in stats['counts'].items():
            if not counts.empty:
                doc.add_heading(f"{category} Distribution", level=3)
                table = doc.add_table(rows=1, cols=2)
                table.rows[0].cells[0].text = category
                table.rows[0].cells[1].text = 'Count'
                
                for item, count in counts.items():
                    row = table.add_row()
                    row.cells[0].text = str(item)
                    row.cells[1].text = str(count)
        
        # Visualizations
        doc.add_heading('Visualizations', level=1)
        plots = self.generate_plots(df)
        
        # Add each plot with its own heading
        plot_titles = {
            'score_dist': 'Score Distribution',
            'type_dist': 'Distribution by Type',
            'genre_dist': 'Genre Distribution',
            'theme_dist': 'Theme Distribution'
        }
        
        for name, fig in plots.items():
            doc.add_heading(plot_titles[name], level=2)
            plot_path = f'temp_{name}.png'
            fig.savefig(plot_path, bbox_inches='tight', dpi=300)
            doc.add_picture(plot_path, width=Inches(6))
            if os.path.exists(plot_path):
                os.remove(plot_path)
        
        # Save document
        doc.save(output_path)
        return True

if __name__ == "__main__":
    analyzer = AnimeDescriptor()
    df = pd.read_csv('Data/AnimeFiltered.csv')
    analyzer.export_report(df, query='Your search query here')